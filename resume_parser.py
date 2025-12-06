"""
Resume Parser Module

Handles extraction of text and structured information from resume files (PDF, DOCX).
Uses regex patterns and spaCy NLP for contact info, skills, experience, and education extraction.
"""
import io
import re
import logging
from typing import Optional, List, Dict, Any, Tuple

import pdfplumber
from docx import Document

try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
except OSError:
    nlp = None
    logging.warning("spaCy model 'en_core_web_sm' not found. Some NLP features will be limited.")

from models import Resume

logger = logging.getLogger(__name__)

# --- Regex Patterns for Contact Information ---
EMAIL_PATTERN = re.compile(
    r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
)

PHONE_PATTERNS = [
    re.compile(r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'),
    re.compile(r'\([0-9]{3}\)\s*[0-9]{3}[-.\s]?[0-9]{4}'),
    re.compile(r'[0-9]{3}[-.\s][0-9]{3}[-.\s][0-9]{4}'),
    # Shorter format: (123) 456-789 (9 digits total)
    re.compile(r'\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{3}'),
]

LINKEDIN_PATTERN = re.compile(
    r'(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9_-]+/?',
    re.IGNORECASE
)

# --- Section Header Patterns ---
# Match section headers anywhere in text (not just line starts)
# For experience, require WORK or PROFESSIONAL prefix to avoid matching "user experience" etc.
EXPERIENCE_HEADERS = re.compile(
    r'\b(WORK\s+EXPERIENCE|PROFESSIONAL\s+EXPERIENCE|EMPLOYMENT(?:\s+HISTORY)?|CAREER\s+HISTORY)\b',
    re.IGNORECASE
)

EDUCATION_HEADERS = re.compile(
    r'\b(EDUCATION(?:AL\s+BACKGROUND)?|ACADEMIC\s+BACKGROUND|QUALIFICATIONS)\b',
    re.IGNORECASE
)

SKILLS_HEADERS = re.compile(
    r'\b(TECHNICAL\s+SKILLS(?:\s+&\s+TOOLS)?|SKILLS|COMPETENCIES|EXPERTISE|TECHNOLOGIES|PROFICIENCIES)\b',
    re.IGNORECASE
)

PROJECTS_HEADERS = re.compile(
    r'\b(PERSONAL\s+PROJECTS|PROJECTS|PORTFOLIO)\b',
    re.IGNORECASE
)

# --- Date Patterns ---
DATE_PATTERN = re.compile(
    r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
    r'[\s,]*\d{4}|'
    r'\d{1,2}/\d{4}|'
    r'\d{4}\s*[-–]\s*(?:\d{4}|present|current)',
    re.IGNORECASE
)

# --- Skills Taxonomy ---
TECH_SKILLS = {
    # Programming Languages
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go', 'golang',
    'rust', 'scala', 'kotlin', 'swift', 'php', 'r', 'matlab', 'perl', 'bash', 'shell',
    # Web Technologies
    'html', 'css', 'react', 'reactjs', 'react.js', 'angular', 'vue', 'vuejs', 'vue.js',
    'node', 'nodejs', 'node.js', 'express', 'django', 'flask', 'fastapi', 'spring',
    'asp.net', '.net', 'rails', 'ruby on rails', 'next.js', 'nextjs', 'nuxt',
    # Databases
    'sql', 'mysql', 'postgresql', 'postgres', 'mongodb', 'redis', 'elasticsearch',
    'cassandra', 'dynamodb', 'oracle', 'sqlite', 'mariadb', 'neo4j', 'graphql',
    # Cloud & DevOps
    'aws', 'amazon web services', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes',
    'k8s', 'terraform', 'ansible', 'jenkins', 'ci/cd', 'github actions', 'gitlab ci',
    'circleci', 'travis ci', 'helm', 'prometheus', 'grafana',
    # Data & ML
    'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'keras', 'scikit-learn',
    'sklearn', 'pandas', 'numpy', 'scipy', 'spark', 'pyspark', 'hadoop', 'airflow',
    'tableau', 'power bi', 'looker', 'data analysis', 'data science', 'nlp',
    'natural language processing', 'computer vision', 'opencv',
    # Tools & Misc
    'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'agile', 'scrum',
    'linux', 'unix', 'windows', 'macos', 'rest', 'restful', 'api', 'microservices',
    'kafka', 'rabbitmq', 'celery', 'nginx', 'apache', 'graphql', 'grpc',
}

SOFT_SKILLS = {
    'leadership', 'communication', 'teamwork', 'problem solving', 'problem-solving',
    'critical thinking', 'project management', 'time management', 'analytical',
    'collaboration', 'adaptability', 'creativity', 'attention to detail',
    'interpersonal', 'presentation', 'negotiation', 'decision making', 'strategic',
    'mentoring', 'coaching', 'cross-functional', 'stakeholder management',
}

# Skill normalization mapping
SKILL_ALIASES = {
    'js': 'javascript',
    'ts': 'typescript',
    'py': 'python',
    'ml': 'machine learning',
    'dl': 'deep learning',
    'ai': 'artificial intelligence',
    'react.js': 'react',
    'reactjs': 'react',
    'vue.js': 'vue',
    'vuejs': 'vue',
    'node.js': 'node',
    'nodejs': 'node',
    'postgres': 'postgresql',
    'k8s': 'kubernetes',
    'sklearn': 'scikit-learn',
    'gcp': 'google cloud',
    'aws': 'amazon web services',
}


# --- PDF Text Extraction ---
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text content from a PDF file.

    Args:
        file_bytes: Raw bytes of the PDF file

    Returns:
        Extracted text as a string, or empty string if extraction fails
    """
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        if "encrypted" in str(e).lower() or "password" in str(e).lower():
            raise ValueError("PDF is encrypted or password-protected")
        raise ValueError(f"Failed to extract text from PDF: {e}")


# --- DOCX Text Extraction ---
def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text content from a DOCX file.

    Args:
        file_bytes: Raw bytes of the DOCX file

    Returns:
        Extracted text as a string
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))

        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}")


# --- Unified Parser Interface ---
def parse_resume(file_bytes: bytes, file_type: str) -> str:
    """
    Parse a resume file and extract text content.

    Args:
        file_bytes: Raw bytes of the file
        file_type: File extension/type ('pdf' or 'docx')

    Returns:
        Cleaned, normalized text ready for further processing
    """
    file_type = file_type.lower().strip('.')

    if file_type == 'pdf':
        raw_text = extract_text_from_pdf(file_bytes)
    elif file_type in ('docx', 'doc'):
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Supported types: pdf, docx")

    # Clean and normalize text
    cleaned = clean_resume_text(raw_text)
    logger.info(f"Extracted {len(cleaned)} characters from {file_type} file")

    return cleaned


def clean_resume_text(text: str) -> str:
    """
    Clean and normalize resume text **while preserving line breaks**.
    This is important so that sections like EDUCATION contain separate
    blocks for each school / degree.
    """
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Collapse spaces/tabs but KEEP newlines
    text = re.sub(r'[ \t]+', ' ', text)

    # Remove trailing spaces before newline
    text = re.sub(r' ?\n', '\n', text)

    # Collapse 3+ blank lines into max 2
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


# --- Contact Information Extraction ---
def extract_email(text: str) -> Optional[str]:
    """Extract email address from text."""
    match = EMAIL_PATTERN.search(text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    """Extract phone number from text."""
    for pattern in PHONE_PATTERNS:
        match = pattern.search(text)
        if match:
            return match.group(0)
    return None


def extract_linkedin(text: str) -> Optional[str]:
    """Extract LinkedIn URL from text."""
    match = LINKEDIN_PATTERN.search(text)
    return match.group(0) if match else None


def extract_name_with_spacy(text: str) -> Optional[str]:
    """Extract person name using spaCy NER."""
    if nlp is None:
        return None

    # Check first 500 chars - name is usually at the top
    doc = nlp(text[:500])
    for ent in doc.ents:
        if ent.label_ == 'PERSON':
            return ent.text
    return None

def extract_name_heuristic(text: str) -> Optional[str]:
    """
    Fallback name extractor when spaCy fails.

    Strategy:
    - Look only at the first few non-empty lines.
    - Prefer a two-word ALL-CAPS token at the very beginning
      (e.g. 'HARPER RUSSO').
    - Otherwise look for 2–4 capitalized tokens that don't look
      like an address or job title line.
    """

    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if not lines:
        return None

    header = lines[:7]

    forbidden_address = {
        "street", "st", "ave", "road", "rd", "suite", "address",
        "city", "zip", "phone", "email", "contact"
    }
    forbidden_titles = {
        "manager", "engineer", "developer", "analyst", "consultant",
        "director", "specialist", "coordinator", "intern",
        "operations", "business", "marketing", "sales", "assistant"
    }

    # 1) FULL-CAPS name at start of line, possibly followed by other text
    #    matches 'HARPER RUSSO' in 'HARPER RUSSO  +123-456-7890 ...'
    caps_name_pattern = re.compile(r'^([A-Z]{2,}(?:\s+[A-Z]{2,}){1,2})\b')

    for line in header:
        m = caps_name_pattern.match(line)
        if m:
            tokens = m.group(1).split()
            if 1 < len(tokens) <= 3:
                return " ".join(t.capitalize() for t in tokens)

    # 2) Normal title-case name candidate
    for line in header:
        lowered = line.lower()

        # Skip address/contact and job-title lines
        if any(w in lowered for w in forbidden_address):
            continue
        if any(w in lowered for w in forbidden_titles):
            continue
        if any(ch.isdigit() for ch in line):
            continue

        parts = line.split()
        if 2 <= len(parts) <= 4 and all(p[0].isalpha() for p in parts):
            # Require that most tokens start with uppercase
            if sum(1 for p in parts if p[0].isupper()) >= len(parts) - 1:
                # Some lines may have a tagline; keep only first 2–3 tokens
                return " ".join(parts[:3])

    return None




def extract_location_with_spacy(text: str) -> Optional[str]:
    """Extract location using spaCy NER."""
    # First try regex for "City, ST" pattern in header (first 500 chars)
    location_pattern = re.compile(
        r'\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)?),\s*([A-Z]{2})\b'
    )
    header_text = text[:500]
    loc_match = location_pattern.search(header_text)
    if loc_match:
        city = loc_match.group(1)
        state = loc_match.group(2)
        # Filter out false positives (degree abbreviations)
        if state not in ('BS', 'BA', 'MS', 'MA', 'MD', 'DO', 'JD', 'PhD'):
            return f"{city}, {state}"

    if nlp is None:
        return None

    # Check first 1000 chars for location
    doc = nlp(text[:1000])
    locations = []
    # Filter out common false positives
    exclude_locations = {'B.S.', 'BS', 'B.A.', 'BA', 'M.S.', 'MS', 'M.A.', 'MA', 'PhD', 'Ph.D.'}
    for ent in doc.ents:
        if ent.label_ in ('GPE', 'LOC') and ent.text not in exclude_locations:
            locations.append(ent.text)

    # Return first unique location found
    return locations[0] if locations else None


# --- Skills Extraction ---
def normalize_skill(skill: str) -> str:
    """Normalize a skill name to its canonical form."""
    skill_lower = skill.lower().strip()
    return SKILL_ALIASES.get(skill_lower, skill_lower)


def extract_skills(text: str) -> List[str]:
    """
    Extract skills from resume text using taxonomy matching.

    Args:
        text: Resume text

    Returns:
        List of unique, normalized skill names
    """
    text_lower = text.lower()
    found_skills = set()

    # Match against tech skills taxonomy
    for skill in TECH_SKILLS:
        # Word boundary matching to avoid partial matches
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower, re.IGNORECASE):
            found_skills.add(normalize_skill(skill))

    # Match against soft skills
    for skill in SOFT_SKILLS:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower, re.IGNORECASE):
            found_skills.add(skill)

    return sorted(list(found_skills))


# --- Section Detection ---
def find_section_boundaries(text: str) -> Dict[str, Tuple[int, int]]:
    """
    Find the start and end positions of major resume sections.

    Returns:
        Dictionary mapping section names to (start, end) positions
    """
    sections = {}
    section_patterns = [
        ('experience', EXPERIENCE_HEADERS),
        ('education', EDUCATION_HEADERS),
        ('skills', SKILLS_HEADERS),
        ('projects', PROJECTS_HEADERS),
    ]

    all_matches = []
    for section_name, pattern in section_patterns:
        for match in pattern.finditer(text):
            all_matches.append((match.start(), section_name, match.end()))

    # Sort by position
    all_matches.sort(key=lambda x: x[0])

    # Determine section boundaries
    for i, (start, name, header_end) in enumerate(all_matches):
        if i + 1 < len(all_matches):
            end = all_matches[i + 1][0]
        else:
            end = len(text)
        # Only keep first occurrence of each section type
        if name not in sections:
            sections[name] = (header_end, end)

    return sections


# --- Experience Extraction ---
def extract_experience(text: str) -> List[Dict[str, Any]]:
    """
    Extract work experience entries from resume text.

    Returns:
        List of experience dictionaries with keys:
        - company: Company name
        - title: Job title
        - dates: Date range string
        - description: List of bullet points/responsibilities
    """
    sections = find_section_boundaries(text)

    if 'experience' not in sections:
        return []

    start, end = sections['experience']
    experience_text = text[start:end]

    experiences = []

    # Date range pattern: "Month Year – Month Year" or "Month Year - Present"
    date_range_pattern = re.compile(
        r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
        r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4})'
        r'\s*[-–]\s*'
        r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
        r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}|Present|Current)',
        re.IGNORECASE
    )

    # Job title patterns - common titles
    title_pattern = re.compile(
        r'\b((?:Senior|Junior|Lead|Staff|Principal|Associate|Assistant)?\s*'
        r'(?:Software|Frontend|Backend|Full[\s-]?Stack|Data|DevOps|Cloud|Machine Learning|ML|AI)?\s*'
        r'(?:Engineer|Developer|Architect|Manager|Analyst|Consultant|Scientist|Director|'
        r'Specialist|Administrator|Designer|Intern|Coordinator|Lead|Superintendent|Supervisor|'
        r'Vice President|VP|CEO|CTO|CFO|COO))\b',
        re.IGNORECASE
    )

    # Find all date ranges in experience section
    date_matches = list(date_range_pattern.finditer(experience_text))

    if date_matches:
        for i, date_match in enumerate(date_matches):
            date_str = date_match.group(0)
            date_start = date_match.start()

            # Get text before this date (company and title info)
            if i == 0:
                before_text = experience_text[:date_start]
            else:
                prev_end = date_matches[i - 1].end()
                before_text = experience_text[prev_end:date_start]

            # Get text after this date (bullet points)
            if i + 1 < len(date_matches):
                after_text = experience_text[date_match.end():date_matches[i + 1].start()]
            else:
                after_text = experience_text[date_match.end():]

            # Extract job title from before text
            title_matches = title_pattern.findall(before_text)
            title = title_matches[-1].strip() if title_matches else None

            # Extract company - try pattern first, then spaCy
            company = None

            # Try pattern: "Company, City, STATE" or "Company, STATE" before job title
            # Look for text that appears before the city/state pattern
            company_loc_pattern = re.compile(
                r'([A-Za-z][A-Za-z\s&\.\-]+?),\s*(?:[A-Za-z\s]+,\s*)?([A-Z]{2})\s',
                re.IGNORECASE
            )
            comp_match = company_loc_pattern.search(before_text)
            if comp_match:
                company = comp_match.group(1).strip()

            # If no company found, try spaCy
            if not company and nlp:
                doc = nlp(before_text[-200:] if len(before_text) > 200 else before_text)
                orgs = [ent.text for ent in doc.ents if ent.label_ == 'ORG']
                if orgs:
                    company = orgs[0]

            # Final fallback: look for capitalized words before the title
            if not company and title:
                fallback_pattern = re.compile(r'([A-Z][A-Za-z\s&]+?)\s+' + re.escape(title), re.IGNORECASE)
                fallback_match = fallback_pattern.search(before_text)
                if fallback_match:
                    company = fallback_match.group(1).strip()

            # Extract bullet points from after_text
            bullets = []
            bullet_pattern = re.compile(r'[•\-\*]\s*(.+?)(?=[•\-\*]|$)')
            bullet_matches = bullet_pattern.findall(after_text)
            bullets = [b.strip() for b in bullet_matches if len(b.strip()) > 10][:5]

            entry = {
                'company': company,
                'title': title,
                'dates': date_str,
                'description': bullets
            }
            experiences.append(entry)

    # Fallback to old behavior if no date ranges found
    elif nlp:
        date_matches = list(DATE_PATTERN.finditer(experience_text))
        doc = nlp(experience_text)
        orgs = [ent.text for ent in doc.ents if ent.label_ == 'ORG']

        for i, date_match in enumerate(date_matches[:3]):  # Limit to 3 entries
            entry = {
                'company': orgs[i] if i < len(orgs) else None,
                'title': None,
                'dates': date_match.group(0),
                'description': []
            }
            experiences.append(entry)

    return experiences


# --- Education Extraction ---
def extract_education(text: str) -> List[Dict[str, Any]]:
    """
    Extract education entries from resume text.

    Strategy:
    1. Prefer the text inside the EDUCATION section if it exists.
    2. First try to find universities/colleges with regex.
    3. If nothing is found (e.g. 'Warner & Spencer'), fall back to splitting
       the education section into blocks separated by blank lines.
    """

    sections = find_section_boundaries(text)

    # ----- 1) Determine which part of the text to scan -----
    if "education" in sections:
        start, end = sections["education"]
        education_text = text[start:end].strip()
    else:
        # No explicit EDUCATION header → scan whole text (we still try to
        # detect education-like blocks using patterns).
        education_text = text
    
    #  Fix: normalize header spacing & bullet separation
    education_text = education_text.replace("•", "\n•")   # bullet varsa satır ayır
    educations: List[Dict[str, Any]] = []

    # ----- 2) First try: universities / colleges regex -----
    university_pattern = re.compile(
        r'\b([A-Z][A-Za-z.&\s-]*(University|College|Institute|Academy|School|Polytechnic))\b',
        re.IGNORECASE
    )

    degree_pattern = re.compile(
        r'(Bachelor|Master|MBA|Bachelors?|Masters?|Diploma|Certificate|Associate|Doctorate)',
        re.IGNORECASE
    )

    field_pattern = re.compile(
        r'in\s+([A-Za-z\s&]+?)(?:,|\||\.|\n|$)', re.IGNORECASE
    )

    date_pattern = re.compile(r'\b(19|20)\d{2}\b')
    gpa_pattern = re.compile(r'GPA[:\s]*([0-4]\.\d{1,2})', re.IGNORECASE)

    uni_matches = list(university_pattern.finditer(education_text))

    if uni_matches:
        # Build entries around each university match
        for match in uni_matches:
            chunk = education_text[match.start(): match.start() + 300]

            degree = None
            for d in degree_pattern.findall(chunk):
                degree = d
                break

            field = None
            for f in field_pattern.finditer(chunk):
                field = f.group(1).strip()
                break

            date = None
            for dt in date_pattern.finditer(chunk):
                date = dt.group(0)
                break

            gpa = None
            g = gpa_pattern.search(chunk)
            if g:
                gpa = g.group(1)

            educations.append({
                "institution": match.group(0).strip(),
                "degree": degree,
                "field": field,
                "dates": date,
                "gpa": gpa,
            })

    # ----- 3) Fallback: block-based parsing (for cases like Warner & Spencer) -----
    if not educations:
        # Split EDUCATION section into blocks separated by blank lines
        # Improved splitting for stacked degrees without blank lines
        blocks = re.split(r'\n(?=[A-Z][A-Za-z& ]+\.?,?$)', education_text.strip())  

        for block in blocks:
            lines = [l.strip() for l in block.splitlines() if l.strip()]
            if not lines:
                continue

            # In education blocks, first line is usually institution,
            # the rest describe degree / field.
            institution = lines[0]
            details = " ".join(lines[1:]) if len(lines) > 1 else ""

            if not institution:
                continue

            degree = None
            for d in degree_pattern.findall(details):
                degree = d
                break

            field = None
            for f in field_pattern.finditer(details):
                field = f.group(1).strip()
                break

            date = None
            for dt in date_pattern.finditer(details):
                date = dt.group(0)
                break

            gpa = None
            g = gpa_pattern.search(details)
            if g:
                gpa = g.group(1)

            educations.append({
                "institution": institution,
                "degree": degree,
                "field": field,
                "dates": date,
                "gpa": gpa,
            })

    return educations


# --- Projects Extraction ---
def extract_projects(text: str) -> List[Dict[str, Any]]:
    """
    Extract project entries from resume text.

    Returns:
        List of project dictionaries with keys:
        - name: Project name
        - technologies: List of technologies used
        - description: List of bullet points describing the project
    """
    sections = find_section_boundaries(text)

    if 'projects' not in sections:
        return []

    start, end = sections['projects']
    projects_text = text[start:end]

    projects = []

    # Pattern for project name with technologies in parentheses
    # e.g., "Festival Playlist Creator for Spotify (Java | JavaScript | React):"
    project_pattern = re.compile(
        r'([A-Z][A-Za-z0-9\s\-]+?)[\s]*\(([^)]+)\)[:\s]*',
        re.IGNORECASE
    )

    # Find all project headers
    project_matches = list(project_pattern.finditer(projects_text))

    if project_matches:
        for i, match in enumerate(project_matches):
            name = match.group(1).strip()
            tech_str = match.group(2)

            # Parse technologies (split by | or ,)
            technologies = [t.strip() for t in re.split(r'[|,]', tech_str) if t.strip()]

            # Get description (bullet points) - text between this project and next
            desc_start = match.end()
            if i + 1 < len(project_matches):
                desc_end = project_matches[i + 1].start()
            else:
                desc_end = len(projects_text)

            desc_text = projects_text[desc_start:desc_end]

            # Extract bullet points
            bullets = []
            bullet_pattern = re.compile(r'[•\-\*]\s*(.+?)(?=[•\-\*]|$)', re.DOTALL)
            for bullet_match in bullet_pattern.finditer(desc_text):
                bullet = bullet_match.group(1).strip()
                # Clean up and limit length
                bullet = re.sub(r'\s+', ' ', bullet)
                if len(bullet) > 20:  # Skip very short bullets
                    bullets.append(bullet[:300])  # Limit bullet length

            project = {
                'name': name,
                'technologies': technologies,
                'description': bullets[:5]  # Limit to 5 bullets per project
            }
            projects.append(project)

    return projects


# --- Summary Extraction ---
def extract_summary(text: str) -> Optional[str]:
    """Extract professional summary/objective from resume."""
    summary_pattern = re.compile(
        r'(?:summary|objective|profile|about\s*me)[:\s]*(.+?)(?=\n\n|experience|education|skills)',
        re.IGNORECASE | re.DOTALL
    )

    match = summary_pattern.search(text[:2000])  # Summary is usually near the top
    if match:
        summary = match.group(1).strip()
        # Limit to reasonable length
        if len(summary) > 50:
            return summary[:500] + '...' if len(summary) > 500 else summary

    return None


# --- Main Orchestrator ---
def extract_resume_details(raw_text: str) -> Resume:
    """
    Extract all structured information from resume text.

    Args:
        raw_text: Raw text extracted from resume file

    Returns:
        Populated Resume object
    """
    extraction_methods = {'regex': [], 'spacy': []}

    # Extract contact information
    email = extract_email(raw_text)
    if email:
        extraction_methods['regex'].append('email')

    phone = extract_phone(raw_text)
    if phone:
        extraction_methods['regex'].append('phone')

    linkedin = extract_linkedin(raw_text)
    if linkedin:
        extraction_methods['regex'].append('linkedin')

    # Full name: try spaCy NER first, then header-based heuristic
    # ----- NAME EXTRACTION -----
    name = extract_name_with_spacy(raw_text)

    # spaCy yakalamazsa header-heuristic kullan (özellikle ALL CAPS için)
    if not name:
        header_text = raw_text.split("\n")[:5]  # sadece ilk satırlar
        name = extract_name_heuristic("\n".join(header_text))
        if name:
            extraction_methods['regex'].append('name_heuristic')

    location = extract_location_with_spacy(raw_text)
    if location:
        extraction_methods['spacy'].append('location')

    # Extract sections
    skills = extract_skills(raw_text)
    if skills:
        extraction_methods['regex'].append('skills')

    experience = extract_experience(raw_text)
    if experience:
        extraction_methods['spacy'].append('experience')

    education = extract_education(raw_text)
    if education:
        extraction_methods['spacy'].append('education')

    projects = extract_projects(raw_text)
    if projects:
        extraction_methods['regex'].append('projects')

    # Build metadata
    metadata = {
        'extraction_methods': extraction_methods,
        'linkedin_url': linkedin,
    }

    return Resume(
        full_name=name,
        email=email,
        phone=phone,
        location=location,
        summary=None,  # Not used for American resumes
        skills=skills,
        experience=experience,
        education=education,
        projects=projects,
        raw_text=raw_text,
        metadata=metadata
    )
